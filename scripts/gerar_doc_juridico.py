# -*- coding: utf-8 -*-
"""Gera o documento .docx para revisão jurídica da Política de Privacidade e Termos de Uso."""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Estilo base
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)

AZUL = RGBColor(0x0A, 0x2A, 0x66)
CINZA = RGBColor(0x60, 0x60, 0x60)


def h1(txt):
    p = doc.add_heading(txt, level=1)
    for r in p.runs:
        r.font.color.rgb = AZUL
    return p


def h2(txt):
    p = doc.add_heading(txt, level=2)
    for r in p.runs:
        r.font.color.rgb = AZUL
    return p


def h3(txt):
    p = doc.add_heading(txt, level=3)
    return p


def para(txt="", bold=False, italic=False, color=None, size=None, align=None):
    p = doc.add_paragraph()
    r = p.add_run(txt)
    r.bold = bold
    r.italic = italic
    if color:
        r.font.color.rgb = color
    if size:
        r.font.size = Pt(size)
    if align:
        p.alignment = align
    return p


def bullets(itens):
    for it in itens:
        p = doc.add_paragraph(style="List Bullet")
        if isinstance(it, tuple):
            r = p.add_run(it[0])
            r.bold = True
            p.add_run(it[1])
        else:
            p.add_run(it)


def checkbox(itens):
    for it in itens:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run("☐  ")
        r.bold = True
        if isinstance(it, tuple):
            rr = p.add_run(it[0])
            rr.bold = True
            p.add_run(it[1])
        else:
            p.add_run(it)


def linha():
    para("―" * 40, color=CINZA, align=WD_ALIGN_PARAGRAPH.CENTER)


# ──────────────────────────────────────────────────────────────────────────
# CAPA
# ──────────────────────────────────────────────────────────────────────────
t = doc.add_heading("CuraDentes", level=0)
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
para("Documento para Revisão Jurídica", bold=True, size=16, align=WD_ALIGN_PARAGRAPH.CENTER, color=AZUL)
para("Política de Privacidade e Termos de Uso · LGPD", size=12, align=WD_ALIGN_PARAGRAPH.CENTER, color=CINZA)
para("")
para("Plataforma: curadentes.com.br", align=WD_ALIGN_PARAGRAPH.CENTER)
para("Versão dos textos analisados: 3 de junho de 2026", align=WD_ALIGN_PARAGRAPH.CENTER, color=CINZA)
para("Revisão 1 (R1) · Documento gerado em: 14 de junho de 2026", align=WD_ALIGN_PARAGRAPH.CENTER, color=CINZA)
doc.add_page_break()

# ──────────────────────────────────────────────────────────────────────────
# 1. OBJETIVO
# ──────────────────────────────────────────────────────────────────────────
h1("1. Objetivo deste documento")
para(
    "Solicitamos a revisão jurídica da Política de Privacidade e dos Termos de Uso da plataforma "
    "CuraDentes, à luz da Lei Geral de Proteção de Dados (LGPD — Lei nº 13.709/2018), do Marco Civil "
    "da Internet (Lei nº 12.965/2014) e do Código de Defesa do Consumidor."
)
para(
    "Os textos já existem e estão publicados no site (seções 5 e 6 deste documento). Pedimos atenção "
    "especial à seção 7 — “Pontos para revisão” —, que reúne lacunas, placeholders e dúvidas "
    "específicas que identificamos internamente e que dependem de orientação jurídica."
)

# ──────────────────────────────────────────────────────────────────────────
# 2. SOBRE A PLATAFORMA
# ──────────────────────────────────────────────────────────────────────────
h1("2. Sobre a plataforma")
para(
    "O CuraDentes é um diretório e motor de busca que conecta pacientes a dentistas. NÃO é uma clínica, "
    "NÃO presta serviços de saúde, NÃO realiza diagnósticos e NÃO intermedia/agenda consultas — atua "
    "apenas como vitrine de profissionais."
)
bullets([
    ("Dois perfis de usuário: ", "paciente (acessa por login social Google) e dentista (“CuraDentes Pro”, cadastro completo)."),
    ("Hospedagem de dados: ", "Supabase (banco PostgreSQL, autenticação e armazenamento de arquivos)."),
    ("Serviços de terceiros: ", "Google (login social), ViaCEP (consulta de CEP), Resend (envio de e-mails transacionais)."),
    ("Modelo de dados: ", "não há venda de dados a terceiros; não há cookies de rastreamento ou publicidade."),
])

# ──────────────────────────────────────────────────────────────────────────
# 3. DADOS PESSOAIS TRATADOS
# ──────────────────────────────────────────────────────────────────────────
h1("3. Dados pessoais tratados")
h3("Dentistas (CuraDentes Pro)")
bullets([
    "Nome completo, e-mail, telefone celular e CPF",
    "CRO (Conselho Regional de Odontologia) e ano de formação",
    "Foto de perfil",
    "Endereço completo do consultório (logradouro, número, bairro, cidade, estado, CEP)",
    "Telefone/WhatsApp do consultório, procedimentos, convênios, formas de pagamento, agenda",
    "Geolocalização (latitude/longitude) derivada do endereço",
])
h3("Pacientes")
bullets([
    "Nome completo e e-mail (via login Google)",
    "Foto de perfil (via Google, quando disponível)",
    "Geolocalização aproximada (latitude/longitude), apenas quando autorizada pelo navegador",
])
h3("Avaliações e dados automáticos")
bullets([
    "Avaliações: nota (1–5), especialidade avaliada e identificador do paciente avaliador",
    "Automáticos: endereço IP, dados de navegação, preferências de busca, cache local (localStorage)",
])

# ──────────────────────────────────────────────────────────────────────────
# 4. MEDIDAS JÁ IMPLEMENTADAS
# ──────────────────────────────────────────────────────────────────────────
h1("4. Medidas de proteção / LGPD já implementadas")
para("Para que o jurídico avalie a suficiência das medidas técnicas e organizacionais já adotadas:")
bullets([
    ("Consentimento explícito: ", "o dentista aceita formalmente na etapa de consentimento do cadastro."),
    ("Direito ao esquecimento: ", "função de exclusão de conta (soft delete) acionável pelo titular."),
    ("Proteção do CPF: ", "o CPF fica isolado em tabela com acesso restrito por linha (RLS); apenas o próprio titular lê seu CPF."),
    ("Geolocalização do paciente — retenção zero: ", "lat/lng é eliminada IMEDIATAMENTE na exclusão da conta (não há retenção da localização precisa após o “esquecimento”)."),
    ("Controle de acesso: ", "Row Level Security (RLS) por usuário no banco de dados."),
    ("Segurança: ", "criptografia em trânsito (HTTPS/TLS); senhas com hash (bcrypt); tokens JWT com expiração."),
    ("Verificação profissional: ", "processo de verificação de CRO dos dentistas."),
    ("Consentimento de comunicações: ", "no cadastro, o dentista opta (opt-in, desmarcado por padrão) por categorias de e-mail opcionais — desempenho do perfil, novidades/dicas e ofertas de parceiros — separado do aceite dos Termos. Pode revogar a qualquer momento por link de descadastro (página pública identificada por token). E-mails essenciais (conta/segurança/serviço) independem do consentimento. As preferências e o token ficam em tabela separada protegida (não expostos pelo SELECT público de curadentespro)."),
])

doc.add_page_break()
# ──────────────────────────────────────────────────────────────────────────
# 5. POLÍTICA DE PRIVACIDADE (texto atual)
# ──────────────────────────────────────────────────────────────────────────
h1("5. Texto ATUAL — Política de Privacidade")
para("Última atualização: 3 de junho de 2026", italic=True, color=CINZA)

h2("1. Introdução")
para("A CuraDentes (“nós”, “nosso” ou “plataforma”) está comprometida com a proteção da privacidade dos seus usuários. Esta Política de Privacidade descreve como coletamos, usamos, armazenamos e compartilhamos suas informações pessoais quando você utiliza nossa plataforma, em conformidade com a Lei Geral de Proteção de Dados Pessoais (LGPD — Lei nº 13.709/2018).")
para("Ao utilizar o CuraDentes, você declara estar ciente e de acordo com as práticas descritas neste documento. Se você é um dentista cadastrado em nossa plataforma (“CuraDentes Pro”), a aceitação formal ocorre na etapa de consentimento do cadastro.")

h2("2. Dados que coletamos")
para("Coletamos as seguintes categorias de dados pessoais:")
para("2.1. Dados dos dentistas (CuraDentes Pro)", bold=True)
bullets([
    "Nome completo, e-mail, telefone celular e CPF",
    "CRO (Conselho Regional de Odontologia) e ano de formação",
    "Foto de perfil (fornecida voluntariamente)",
    "Dados de endereço: logradouro, número, complemento, bairro, cidade, estado e CEP",
    "Telefone e WhatsApp do consultório",
    "Procedimentos realizados, convênios aceitos e formas de pagamento",
    "Agenda de horários de atendimento",
    "Política de cancelamento e observações",
    "Geolocalização (latitude/longitude) gerada automaticamente a partir do endereço",
])
para("2.2. Dados dos pacientes", bold=True)
bullets([
    "Nome completo e e-mail (fornecidos pelo login do Google)",
    "Foto de perfil (fornecida pelo Google, quando disponível)",
    "Localização aproximada (latitude/longitude) quando autorizada pelo navegador",
])
para("2.3. Dados de avaliações", bold=True)
bullets([
    "Nota (1 a 5) atribuída a um dentista",
    "Especialidade avaliada",
    "Registro do paciente avaliador (ID anônimo)",
])
para("2.4. Dados coletados automaticamente", bold=True)
bullets([
    "Endereço IP e dados de navegação (user agent, páginas acessadas)",
    "Preferências de busca e filtros utilizados",
    "Dados de cache armazenados localmente no navegador (localStorage)",
])

h2("3. Finalidade do tratamento")
para("Utilizamos seus dados para as seguintes finalidades:")
bullets([
    ("Viabilizar a plataforma: ", "conectar pacientes a dentistas, exibir perfis profissionais, permitir busca por localização e especialidade."),
    ("Gerenciar sua conta: ", "autenticação, recuperação de senha e manutenção do cadastro profissional."),
    ("Recomendações: ", "exibir dentistas próximos com base na sua localização ou filtros selecionados."),
    ("Melhoria do serviço: ", "analisar padrões de uso para aprimorar a experiência e corrigir falhas."),
    ("Comunicação: ", "enviar comunicados operacionais (confirmação de cadastro, alterações na plataforma)."),
    ("Segurança: ", "prevenir fraudes, garantir a integridade dos dados e cumprir obrigações legais."),
])

h2("4. Base legal para o tratamento")
para("Em conformidade com o art. 7º da LGPD, as bases legais que utilizamos são:")
bullets([
    ("Consentimento (art. 7º, I): ", "para o cadastro de dentistas, mediante aceite expresso na etapa de consentimento."),
    ("Execução de contrato (art. 7º, V): ", "para viabilizar o funcionamento da plataforma e a conexão entre pacientes e dentistas."),
    ("Legítimo interesse (art. 7º, IX): ", "para análises internas, melhoria do serviço e segurança da plataforma."),
    ("Obrigação legal (art. 7º, II): ", "para cumprir requisitos regulatórios e ordens judiciais."),
])

h2("5. Compartilhamento de dados")
para("Compartilhamos seus dados apenas nas seguintes circunstâncias:")
bullets([
    ("Com prestadores de serviço: ", "como o Supabase (hospedagem de banco de dados e autenticação) e ViaCEP (consulta de CEP). Esses parceiros são contratualmente obrigados a proteger seus dados."),
    ("Com outros usuários: ", "os dados do perfil do dentista (nome, CRO, endereços, agenda, convênios) são públicos para pacientes que utilizam a busca."),
    ("Por determinação legal: ", "mediante ordem judicial ou requisição de autoridade competente."),
])
para("O CuraDentes não vende dados pessoais para terceiros.", italic=True, color=CINZA)

h2("6. Cookies e tecnologias semelhantes")
para("Utilizamos o localStorage do navegador para armazenar preferências de navegação e dados de cache, como:")
bullets([
    "Rascunho de cadastro (para evitar perda de dados durante o registro)",
    "Cache de consultas de CEP (para evitar reconsultas desnecessárias à API ViaCEP)",
    "Cache de localização do usuário (para lembrar sua cidade por 30 minutos)",
    "Cache de resultados de busca (para exibição offline-first)",
])
para("Atualmente não utilizamos cookies de rastreamento ou publicidade. Caso isso mude, atualizaremos esta política com antecedência.")

h2("7. Seus direitos (LGPD)")
para("Nos termos da LGPD (arts. 17 a 22), você possui os seguintes direitos em relação aos seus dados pessoais:")
bullets([
    ("Confirmação e acesso: ", "saber se tratamos seus dados e acessá-los"),
    ("Correção: ", "corrigir dados incompletos, inexatos ou desatualizados"),
    ("Anonimização ou exclusão: ", "solicitar a eliminação de dados desnecessários"),
    ("Portabilidade: ", "solicitar a transferência a outro fornecedor"),
    ("Revogação do consentimento: ", "retirar seu consentimento a qualquer momento"),
    ("Oposição: ", "opor-se ao tratamento com base no legítimo interesse"),
    ("Informação sobre compartilhamento: ", "saber com quais entidades compartilhamos seus dados"),
])
para("Para exercer seus direitos, entre em contato pelo e-mail/formulário da seção “Contato”. Responderemos em até 15 dias.")

h2("8. Prazo de retenção")
para("Armazenamos seus dados pessoais pelo período necessário para cumprir as finalidades descritas nesta política:")
bullets([
    ("Dentistas: ", "enquanto a conta permanecer ativa. Após exclusão da conta, os dados serão anonimizados ou eliminados em até 90 dias, exceto quando houver obrigação legal de retenção."),
    ("Pacientes: ", "enquanto o usuário mantiver seu vínculo com a plataforma (login social). Ao excluir a conta, sua geolocalização (latitude/longitude) é eliminada imediatamente — não há retenção da localização precisa após o “esquecimento”. Avaliações anônimas podem ser mantidas após o desligamento."),
    ("Dados de cache local: ", "por até 7 dias (CEP) ou 30 minutos (localização), ou até a limpeza manual do navegador."),
])

h2("9. Segurança dos dados")
para("Adotamos medidas técnicas e organizacionais para proteger seus dados pessoais:")
bullets([
    "Criptografia em trânsito via HTTPS/TLS",
    "Senhas armazenadas com hash (bcrypt) pelo Supabase Auth",
    "Políticas de acesso baseadas em função (RLS — Row Level Security) no banco de dados",
    "Separação lógica dos dados entre diferentes usuários",
    "Token de autenticação JWT com expiração",
    "Práticas de desenvolvimento seguro (sanitização de inputs, proteção contra XSS e CSRF)",
])

h2("10. Alterações nesta política")
para("Esta política pode ser atualizada periodicamente para refletir mudanças em nossas práticas ou na legislação. A data da última atualização está indicada no início do documento.")
para("Notificaremos os dentistas cadastrados sobre alterações significativas por e-mail. Recomendamos que todos os usuários revisem esta página regularmente.")

h2("11. Contato")
para("Para exercer seus direitos, esclarecer dúvidas ou reportar preocupações sobre esta política, entre em contato conosco:")
bullets([
    ("E-mail: ", "privacidade@curadentes.com.br"),
    ("Endereço: ", "[Inserir endereço da empresa]  ← PLACEHOLDER A PREENCHER"),
])
para("Responderemos em até 15 dias úteis.")

doc.add_page_break()
# ──────────────────────────────────────────────────────────────────────────
# 6. TERMOS DE USO (texto atual)
# ──────────────────────────────────────────────────────────────────────────
h1("6. Texto ATUAL — Termos de Uso")
para("Última atualização: 3 de junho de 2026", italic=True, color=CINZA)

h2("1. Aceitação dos Termos")
para("Ao utilizar o CuraDentes (“plataforma”), você declara ter lido, compreendido e concordado com estes Termos de Uso. Se você não concorda com qualquer parte destes termos, não utilize a plataforma.")
para("Para dentistas (“CuraDentes Pro”), a aceitação formal ocorre na etapa de consentimento do cadastro. Para pacientes, o aceite ocorre no momento do login.")

h2("2. Definições")
bullets([
    ("Plataforma: ", "o site CuraDentes e todos os seus serviços, acessível através do endereço curadentes.com.br"),
    ("Paciente: ", "usuário que busca dentistas para agendamento de consultas"),
    ("Dentista / Profissional: ", "profissional de odontologia cadastrado no CuraDentes Pro"),
    ("CuraDentes Pro: ", "plano de cadastro profissional que permite ao dentista gerenciar seu perfil e ser encontrado por pacientes"),
    ("LGPD: ", "Lei Geral de Proteção de Dados Pessoais (Lei nº 13.709/2018)"),
])

h2("3. Cadastro do paciente")
para("Pacientes podem acessar a plataforma por meio de login social (Google). Ao fazê-lo, autorizam o CuraDentes a coletar nome, e-mail e foto do perfil do Google para fins de identificação e funcionamento da plataforma.")
para("O paciente é responsável pela veracidade das informações fornecidas e pela guarda de suas credenciais de acesso.")

h2("4. Cadastro do dentista (CuraDentes Pro)")
para("O cadastro de dentista é realizado em etapas e requer o preenchimento de dados pessoais e profissionais, conforme descrito na Política de Privacidade.")
para("4.1. Veracidade das informações", bold=True)
para("O dentista declara, sob as penas da lei, que todas as informações fornecidas são verdadeiras, especialmente o CRO. A plataforma se reserva o direito de verificar a autenticidade dos dados e suspender contas que contenham informações falsas.")
para("4.2. Responsabilidade pelo perfil", bold=True)
para("O dentista é o único responsável pelo conteúdo publicado em seu perfil, incluindo textos, horários, convênios e formas de pagamento. A plataforma não se responsabiliza por informações desatualizadas ou incorretas.")
para("4.3. Exclusividade de cadastro", bold=True)
para("Cada dentista pode possuir apenas uma conta ativa. Cadastros duplicados poderão ser removidos sem aviso prévio.")

h2("5. Uso da plataforma")
para("A plataforma CuraDentes conecta pacientes a dentistas. Não realizamos agendamentos, não intermediamos consultas e não somos responsáveis pela qualidade do serviço prestado pelos profissionais cadastrados.")
para("5.1. Papel do CuraDentes", bold=True)
para("O CuraDentes atua exclusivamente como um diretório e motor de busca de profissionais de odontologia. Não somos uma clínica odontológica, não prestamos serviços de saúde e não realizamos diagnósticos.")
para("5.2. Avaliações", bold=True)
para("Pacientes podem avaliar dentistas com base em suas experiências. As avaliações refletem a opinião pessoal do paciente e não representam a posição do CuraDentes. Reservamo-nos o direito de remover avaliações que violem estes termos (discurso de ódio, informações falsas, spam).")

h2("6. Obrigações dos usuários")
para("Ao utilizar a plataforma, você concorda em:")
bullets([
    "Não fornecer informações falsas ou fraudulentas",
    "Não utilizar a plataforma para fins ilegais ou não autorizados",
    "Não reproduzir, distribuir ou modificar o conteúdo da plataforma sem autorização",
    "Não coletar dados de outros usuários de forma automatizada (scraping)",
    "Não publicar conteúdo ofensivo, difamatório, discriminatório ou que viole direitos de terceiros",
    "Não tentar acessar áreas restritas da plataforma sem autorização",
    "Manter seus dados de cadastro atualizados",
])

h2("7. Responsabilidade")
para("7.1. Limitação de responsabilidade", bold=True)
para("O CuraDentes não se responsabiliza por danos diretos ou indiretos decorrentes de:")
bullets([
    "Atendimento odontológico prestado pelos profissionais cadastrados",
    "Informações desatualizadas ou incorretas nos perfis dos dentistas",
    "Indisponibilidade temporária da plataforma por manutenção ou falhas técnicas",
    "Perda de dados decorrente de ação do usuário (limpeza de cache, exclusão de conta)",
    "Atos de terceiros (hackers, malware) desde que adotadas as medidas de segurança cabíveis",
])
para("7.2. Responsabilidade do dentista", bold=True)
para("O dentista é civil e criminalmente responsável pelos serviços odontológicos que presta. O CuraDentes não exerce qualquer controle ou supervisão sobre a atividade profissional dos dentistas cadastrados.")
para("7.3. Indenização", bold=True)
para("O usuário concorda em indenizar o CuraDentes por quaisquer perdas, danos ou custos decorrentes do uso inadequado da plataforma ou da violação destes termos.")

h2("8. Propriedade intelectual")
para("Todos os direitos de propriedade intelectual sobre a plataforma CuraDentes, incluindo código-fonte, design, logotipos, marcas e conteúdo editorial, são de propriedade exclusiva do CuraDentes.")
para("O conteúdo cadastrado pelos dentistas (fotos, textos, horários) é de propriedade do respectivo profissional, que concede ao CuraDentes uma licença não exclusiva para exibi-lo na plataforma.")

h2("9. Cancelamento e suspensão")
para("9.1. Pelo usuário", bold=True)
para("O dentista pode solicitar a exclusão de sua conta a qualquer momento, entrando em contato conosco. Os dados serão anonimizados ou eliminados conforme a Política de Privacidade.")
para("9.2. Pela plataforma", bold=True)
para("O CuraDentes se reserva o direito de suspender ou encerrar contas que violem estes Termos de Uso, sem prejuízo das demais medidas legais cabíveis.")

h2("10. Disposições gerais")
para("10.1. Legislação aplicável", bold=True)
para("Estes Termos de Uso são regidos pela legislação brasileira, em especial pelo Código Civil (Lei nº 10.406/2002) e pelo Marco Civil da Internet (Lei nº 12.965/2014).")
para("10.2. Foro", bold=True)
para("Fica eleito o foro da comarca de [cidade/estado] para dirimir quaisquer controvérsias decorrentes destes Termos, com renúncia a qualquer outro foro, por mais privilegiado que seja.  ← PLACEHOLDER A PREENCHER")
para("10.3. Alterações", bold=True)
para("Estes Termos podem ser alterados a qualquer momento. Dentistas cadastrados serão notificados por e-mail sobre alterações significativas. O uso continuado da plataforma após as alterações constitui aceitação dos novos termos.")
para("10.4. Nulidade parcial", bold=True)
para("Se qualquer cláusula destes Termos for considerada inválida ou inexequível, as demais cláusulas permanecerão em pleno vigor.")

h2("11. Contato")
para("Para dúvidas, reclamações ou notificações relacionadas a estes Termos de Uso:")
bullets([("E-mail: ", "contato@curadentes.com.br")])

doc.add_page_break()
# ──────────────────────────────────────────────────────────────────────────
# 7. PONTOS PARA REVISÃO
# ──────────────────────────────────────────────────────────────────────────
h1("7. Pontos específicos para revisão do advogado")
para("Itens que identificamos internamente e que dependem de orientação/decisão jurídica. Os dois primeiros são placeholders que precisam ser preenchidos.", italic=True, color=CINZA)

h3("Placeholders a preencher")
checkbox([
    ("Endereço da empresa — ", "a Política de Privacidade (seção 11) traz “[Inserir endereço da empresa]”. Definir o endereço cadastral."),
    ("Foro — ", "os Termos (cláusula 10.2) trazem “comarca de [cidade/estado]”. Definir a comarca de eleição."),
])

h3("Encarregado / DPO e governança")
checkbox([
    ("Encarregado pelo tratamento de dados (DPO, art. 41 da LGPD) — ", "não há indicação de encarregado nos textos. Avaliar a obrigatoriedade para este porte/atividade e, se cabível, incluir nome e canal de contato."),
    ("Canal de exercício de direitos — ", "confirmar se privacidade@curadentes.com.br está ativo e operacionalizado."),
])

h3("Retenção e exclusão")
checkbox([
    ("Prazo de 90 dias (dentistas) — ", "a Política prevê eliminação “em até 90 dias” após a exclusão da conta. Validar a adequação do prazo e quais dados têm obrigação legal de retenção (ex.: fiscais/CRO)."),
    ("Avaliações “anônimas” mantidas após o desligamento — ", "o sistema guarda o identificador do paciente avaliador. Avaliar se podem ser tratadas como anônimas e a base legal para mantê-las."),
])

h3("Consentimento e categorias especiais")
checkbox([
    ("Aceite do paciente “no momento do login” — ", "avaliar se constitui consentimento válido e registrável (prova do aceite), ou se exige um aceite explícito adicional."),
    ("Menores de idade — ", "a odontologia atende crianças/adolescentes; os textos não tratam de dados de menores (art. 14 da LGPD). Avaliar necessidade de cláusula e de consentimento parental."),
    ("Necessidade do CPF do dentista — ", "avaliar finalidade/minimização: o CPF é necessário? Em caso afirmativo, deixar a finalidade explícita."),
])

h3("Consentimento de comunicações (e-mail marketing)")
checkbox([
    ("Opt-in granular — ", "validar o consentimento por categoria (desempenho, novidades, ofertas de parceiros), desmarcado por padrão e separado do aceite dos Termos."),
    ("Base legal essenciais x opcionais — ", "confirmar: e-mails essenciais por execução de contrato; e-mails opcionais por consentimento (revogável)."),
    ("Descadastro — ", "validar o mecanismo de cancelamento por token (link nos e-mails + página pública /descadastro) e o registro de data/hora do consentimento e da revogação."),
    ("Ofertas de parceiros — ", "confirmar transparência sobre com quais parceiros há comunicação/compartilhamento e se exige consentimento destacado."),
])

h3("Transferência internacional e terceiros")
checkbox([
    ("Transferência internacional de dados — ", "Supabase e Google podem armazenar/processar dados fora do Brasil. Avaliar a necessidade de cláusula de transferência internacional (arts. 33–36 da LGPD) e a localização (região) do banco."),
    ("Operadores/Sub-operadores — ", "confirmar contratos de tratamento (DPA) com Supabase, Google, ViaCEP e Resend (provedor de e-mail, não citado nos textos)."),
])

h3("Consumidor e responsabilidade")
checkbox([
    ("Cláusulas de limitação de responsabilidade e indenização (Termos 7) — ", "revisar à luz do CDC (vedação a cláusulas abusivas em relação de consumo)."),
    ("Responsabilidade por avaliações (difamação) — ", "revisar Termos 5.2 frente ao Marco Civil (responsabilidade de provedor por conteúdo de terceiros)."),
])

h3("Inconsistências de texto a padronizar")
checkbox([
    ("Prazo de resposta — ", "a Política diz “15 dias” (seção 7) e “15 dias úteis” (seção 11). Padronizar."),
    ("“ID anônimo” do avaliador — ", "a seção 2.3 chama de “anônimo”, mas o identificador vincula a um paciente. Ajustar a redação conforme orientação."),
    ("Data de “última atualização” — ", "atualizar (3 de junho de 2026) quando os textos forem revisados/aprovados."),
])

doc.add_page_break()
# ──────────────────────────────────────────────────────────────────────────
# 8. IDENTIDADE VISUAL — FONTES E TIPOGRAFIA
# ──────────────────────────────────────────────────────────────────────────
h1("8. Identidade visual — fontes e tipografia")
para(
    "A marca utiliza uma fonte própria, a “CuraDentes”, criada a partir de uma fonte "
    "de licença aberta. Esta seção descreve a origem e pede validação da conformidade "
    "de licenciamento, já que envolve licença de software de fonte e a fonte comercial Nexa."
)

h3("Fonte “CuraDentes” (da marca)")
bullets([
    ("Origem: ", "derivada da fonte Montserrat, licenciada sob a SIL Open Font License 1.1 (OFL) — licença que permite expressamente usar, modificar, renomear e redistribuir, inclusive para fins comerciais."),
    ("Estilo: ", "letras geométricas (mesmo “espírito” visual da fonte Nexa), com os numerais da Montserrat."),
    ("Conformidade aplicada: ", "mantido o aviso de copyright original da Montserrat; removido o nome reservado “Montserrat” dos campos de nome do arquivo; a licença OFL acompanha os arquivos da fonte (arquivo OFL-CuraDentes.txt)."),
    ("Nexa NÃO usada em produção: ", "a fonte comercial Nexa (Fontfabric, versão trial) foi utilizada apenas para comparação visual interna; nenhum arquivo da Nexa integra o material de produção (e-mail/site)."),
])

h3("Pontos para revisão (fontes / identidade)")
checkbox([
    ("Conformidade com a SIL OFL 1.1 — ", "validar a renomeação e a redistribuição da fonte derivada (retenção do aviso de copyright, não uso do nome reservado “Montserrat”, distribuição da licença OFL junto dos arquivos)."),
    ("Não infração da Nexa — ", "confirmar que, por não usarmos arquivos da Nexa e por o desenho de tipos (a aparência das letras) em regra não ser protegido por direito autoral no Brasil, não há violação à Fontfabric. Garantir que a versão trial da Nexa não seja distribuída nem usada em produção."),
    ("Titularidade dos logotipos — ", "verificar a propriedade/cessão de direitos dos logotipos (arquivos PNG da marca): quem os criou e se há instrumento de cessão por escrito."),
    ("Registro de marca — ", "avaliar o registro das marcas “CuraDentes” e “CuraDentes Pro” (nominativa e/ou figurativa, incluindo o nome da fonte) junto ao INPI."),
])

doc.add_page_break()
# ──────────────────────────────────────────────────────────────────────────
# ANEXO A — AMOSTRAS VISUAIS (renderizadas com as fontes reais)
# ──────────────────────────────────────────────────────────────────────────
h1("Anexo A — Amostras visuais (Nexa × CuraDentes)")
para("Comparação visual entre a fonte comercial Nexa (exibida apenas para avaliação) e a fonte própria CuraDentes (derivada da Montserrat, SIL OFL 1.1) usada em produção.")

import io as _io
import zipfile as _zip
from PIL import Image as _Img, ImageDraw as _Draw, ImageFont as _Font
from docx.shared import Inches as _In

_PROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
_EFONTS = os.path.join(_PROOT, "email marketing - informativo", "plano mensal", "fonts")
_ZIP = os.path.join(_PROOT, "email marketing - informativo", "plano mensal", "Nexa-Trial-Book.zip")
_NAVY, _PINK, _MUT = (11, 28, 48), (0xB5, 0x00, 0x48), (0x8E, 0x8E, 0x93)


def _nexa(sz, bold=True):
    z = _zip.ZipFile(_ZIP)
    return _Font.truetype(_io.BytesIO(z.read("Nexa-Trial-Bold.ttf" if bold else "Nexa-Trial-Light.ttf")), sz)


def _cd(sz, bold=True):
    return _Font.truetype(os.path.join(_EFONTS, "CuraDentes-Bold.ttf" if bold else "CuraDentes-Light.ttf"), sz)


def _lockup(d, x, y, fb, fl, pink):
    for t, f, c in [("Cura", fb, _NAVY), ("Dentes", fl, _NAVY), (" Pro", fb, pink)]:
        d.text((x, y), t, font=f, fill=c)
        x += d.textlength(t, font=f)


def _emb(img):
    bio = _io.BytesIO()
    img.save(bio, "PNG")
    bio.seek(0)
    doc.add_picture(bio, width=_In(6.3))


# A.1 e A.2 — amostras visuais (dependem da Nexa, trial não versionada)
if os.path.exists(_ZIP):
    para("A.1. Impacto na logo — wordmark “CuraDentes Pro” (Cura Bold + Dentes Light)", bold=True)
    img = _Img.new("RGB", (1600, 540), "white")
    d = _Draw.Draw(img)
    lbl = _cd(30, True)
    d.text((50, 30), "NEXA (original):", font=lbl, fill=_MUT)
    _lockup(d, 50, 80, _nexa(100, True), _nexa(100, False), _PINK)
    d.text((50, 310), "CURADENTES (nossa):", font=lbl, fill=_PINK)
    _lockup(d, 50, 360, _cd(100, True), _cd(100, False), _PINK)
    _emb(img)

    para("A.2. Numerais (CuraDentes usa os da Montserrat)", bold=True)
    img2 = _Img.new("RGB", (1600, 430), "white")
    d2 = _Draw.Draw(img2)
    d2.text((50, 25), "NEXA:", font=_cd(30, True), fill=_MUT)
    d2.text((50, 70), "0 1 2 3 4 5 6 7 8 9", font=_nexa(86, True), fill=_NAVY)
    d2.text((50, 235), "CURADENTES:", font=_cd(30, True), fill=_PINK)
    d2.text((50, 280), "0 1 2 3 4 5 6 7 8 9", font=_cd(86, True), fill=_NAVY)
    _emb(img2)
else:
    para("As amostras visuais comparativas (Nexa × CuraDentes) não foram embutidas neste ambiente "
         "porque o arquivo da fonte Nexa (versão trial, não versionada) não está disponível. A "
         "comparação visual completa permanece no arquivo “comparativo-fontes.html”.",
         italic=True, size=10, color=CINZA)

para("Versão interativa completa (alfabeto, pesos, sobreposição e texto corrido) no arquivo "
     "“comparativo-fontes.html”.", italic=True, size=10, color=CINZA)

para("")
linha()
para("Documento preparado para fins de revisão jurídica. Os textos das seções 5 e 6 reproduzem fielmente o conteúdo publicado em curadentes.com.br/privacidade e curadentes.com.br/termos.", italic=True, size=9, color=CINZA)

import os
out = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "CuraDentes_Revisao_Juridica-R1.docx")
doc.save(out)
print("OK:", out)
